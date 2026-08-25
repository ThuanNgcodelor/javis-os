import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from pathlib import Path

doc = docx.Document()

# Set standard margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Title
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_title = p_title.add_run("TÀI LIỆU HƯỚNG DẪN & BÁO CÁO HỆ THỐNG CFC AI\nCHATBOT CONTROL CENTER v2.1")
run_title.font.size = Pt(20)
run_title.font.bold = True
run_title.font.color.rgb = RGBColor(79, 70, 229)

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_sub = p_sub.add_run("Dự án Chatbot Thông Minh Đa Kênh — ZeO Vietnam & CFC Cò Bay\nNgày cập nhật: 14/08/2026")
run_sub.font.size = Pt(11)
run_sub.font.italic = True
run_sub.font.color.rgb = RGBColor(100, 116, 139)

doc.add_paragraph()

# Heading 1
def add_h1(text):
    h = doc.add_heading(text, level=1)
    h.paragraph_format.space_before = Pt(16)
    h.paragraph_format.space_after = Pt(6)
    for r in h.runs:
        r.font.size = Pt(14)
        r.font.bold = True
        r.font.color.rgb = RGBColor(30, 41, 59)
    return h

def add_h2(text):
    h = doc.add_heading(text, level=2)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(4)
    for r in h.runs:
        r.font.size = Pt(12)
        r.font.bold = True
        r.font.color.rgb = RGBColor(79, 70, 229)
    return h

def add_bullet(text, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        r_b = p.add_run(bold_prefix + " ")
        r_b.bold = True
    p.add_run(text)

# --- PHẦN 1 ---
add_h1("1. TỔNG QUAN HỆ THỐNG VÀ KIẾN TRÚC CÔNG NGHỆ")
p = doc.add_paragraph(
    "Hệ thống CFC AI Chatbot Control Center là nền tảng quản trị và hỗ trợ khách hàng đa kênh "
    "(Messenger, Web, Telegram) cho hai thương hiệu chính: ZeO Vietnam (Hóa mỹ phẩm gia dụng) và "
    "CFC Cò Bay (Phân bón & Dinh dưỡng cây trồng Cần Thơ). Hệ thống kết hợp giữa RAG Semantic Search "
    "chính xác tốc độ cao và AI thế hệ mới để tự động hóa tối đa quy trình chăm sóc khách hàng và chốt đơn."
)
p.paragraph_format.space_after = Pt(8)

add_h2("Các Công Nghệ Cốt Lõi Được Tích Hợp:")
add_bullet("FastAPI (Python 3.9+) làm backend API server bất đồng bộ hiệu năng cao.", "• Backend Service:")
add_bullet("RediSearch Vector Search (KNN HNSW/FLAT, BGE-M3 1024 dim) lưu trữ Knowledge Base và Session khách.", "• Cơ sở dữ liệu Vector:")
add_bullet("Ollama Local (Mô hình BGE-M3 tạo embedding tiếng Việt cực chuẩn và Qwen 2.5:7b viết lại câu tự nhiên).", "• AI Xử Lý Cục Bộ:")
add_bullet("Google Gemini 2.0 Flash (Free), OpenRouter, Groq tự động fallback khi phân tích tài liệu và viết báo cáo.", "• AI Đám Mây Miễn Phí:")
add_bullet("Tự động bắn thông báo Lead có SĐT, Báo cáo điều hành và cảnh báo câu hỏi cần hỗ trợ tức thì.", "• Telegram Notifier:")
add_bullet("Khớp chính xác tên sản phẩm, quy cách, từ khóa (kể cả không dấu) và gửi link mua hàng đúng chuẩn.", "• Shopee Matcher:")

# --- PHẦN 2 ---
add_h1("2. DANH SÁCH CÁC TÍNH NĂNG ĐÃ HOÀN THIỆN ĐẦY ĐỦ")

add_h2("Nhóm A — Cấp Thiết:")
add_bullet("Cho phép kéo thả hoặc bấm tải lên file tài liệu .md/.txt trực tiếp từ trình duyệt. Hệ thống tự động phân đoạn (chunking), tạo embedding và lưu vào Redis ngay.", "1. Upload File .md:")
add_bullet("Nhập đường dẫn Google Sheets (chế độ chia sẻ công khai). Hệ thống tự động đọc bảng dữ liệu FAQ (cột Câu hỏi | Câu trả lời) và nạp thẳng vào Vector Index.", "2. Import Google Sheets:")
add_bullet("Giao diện quản trị danh mục sản phẩm Shopee trực quan. Thêm mới, chỉnh sửa giá, khuyến mãi, quy cách, từ khóa nhận diện và link Shopee.", "3. Shopee Catalog CRUD:")
add_bullet("Background worker tự động chạy ngầm mỗi 10 phút để tải bảng giá mới nhất từ Google Sheets của team Sale và cập nhật vào Catalog.", "4. Auto-Sync Shopee 10 phút:")
add_bullet("Score >= 78% (Tự tin trả lời trực tiếp) | Score >= 55% (Trả lời + Viết lại bằng AI) | Score < 55% (Tự động gửi lời xin lỗi lịch sự, bắn thông báo Telegram cho Admin và đẩy câu hỏi vào Learning Queue).", "5. Fallback Thông Minh 3 Lớp:")

add_h2("Nhóm B — Quản Lý Hội Thoại & Khách Hàng:")
add_bullet("Nhấn vào khách hàng để xem toàn bộ dòng thời gian tin nhắn của khách và phản hồi của bot, kèm nhãn intent và thời gian.", "1. Xem Lịch Sử Chat Đầy Đủ:")
add_bullet("Lọc nhanh khách hàng theo: Đã có SĐT / Chưa có SĐT, Lọc theo Lead Stage (Mới, Đang thu thập, Lead sẵn, Qualified, Chuyển admin, Đã xử lý).", "2. Bộ Lọc Nâng Cao:")
add_bullet("Xuất danh sách khách hàng ra file .CSV tải về máy tính để bàn giao cho đội ngũ telesale / marketing.", "3. Xuất Dữ Liệu CSV:")
add_bullet("Admin có thể ghi chú riêng tư về từng khách (nhu cầu, lịch hẹn gọi lại) và gắn tag phân loại (HOT LEAD, CHỜ BÁO GIÁ, ĐÃ CHỐT...).", "4. Ghi Chú Admin & Tags:")

add_h2("Nhóm C — AI & Phân Tích Thông Minh:")
add_bullet("AI tự động quét toàn bộ các câu hỏi bot chưa tự tin trả lời trong Learning Queue, gom nhóm các câu hỏi cùng chủ đề, đề xuất tên intent và câu trả lời chuẩn. Admin chỉ cần 1 click để duyệt vào FAQ.", "1. AI Tự Đề Xuất FAQ:")
add_bullet("Biểu đồ cột trực quan theo dõi lượng khách hàng mới và tỷ lệ thu thập SĐT trong 7 ngày gần nhất ngay trên Dashboard.", "2. Biểu Đồ Xu Hướng 7 Ngày:")
add_bullet("Tạo báo cáo tổng hợp tình hình kinh doanh, số lượng khách, top intent trong ngày bằng AI và gửi nhanh qua Telegram.", "3. Báo Cáo Điều Hành AI:")

# --- PHẦN 3 ---
add_h1("3. CẤU TRÚC THƯ MỤC SOURCE CODE")
p_tree = doc.add_paragraph()
tree_text = """javis/
├── server/
│   ├── main.py                  # FastAPI app chính + Background Workers (Shopee 10m, Snapshot 1h)
│   ├── admin_routes.py          # Toàn bộ API Endpoints cho Admin Dashboard
│   ├── rag_search.py            # Tìm kiếm ngữ nghĩa + 3-Layer Fallback
│   ├── embedder.py              # Xử lý embedding qua Ollama BGE-M3 (1024 dims)
│   ├── knowledge_sync.py        # Đồng bộ FAQ snapshot vào RediSearch Vector Index
│   ├── document_ingestor.py     # Nạp và phân đoạn tài liệu .md vào Vector Index
│   ├── shopee_matcher.py        # Module nhận diện và khớp sản phẩm Shopee
│   ├── telegram_notifier.py     # Gửi thông báo Lead, Cảnh báo Fallback, Báo cáo qua Telegram
│   ├── ai_engine.py             # Xử lý Cloud AI (Gemini, OpenRouter, Groq)
│   ├── ai_reporter.py           # Sinh Báo Cáo Điều Hành Kinh Doanh Hàng Ngày
│   ├── settings.json            # File lưu toàn bộ cấu hình API Keys & Ngưỡng RAG
│   └── static/                  # Giao diện Admin Dashboard tách đa file
│       ├── admin.html           # Shell HTML gọn sạch, bố cục Sidebar, Topbar, Modals
│       ├── css/
│       │   ├── base.css         # Reset CSS, Biến màu Dark Mode, Typography
│       │   ├── layout.css       # Bố cục Sidebar nhóm, Topbar Breadcrumb, Footer
│       │   └── components.css   # Cards, Tables, Buttons, Badges, Modals, Chat Timeline, Trend Charts
│       └── js/
│           ├── core.js          # Helpers dùng chung, Navigation, Quản lý State
│           └── pages/
│               ├── dashboard.js # Dashboard Stats & Trend Analytics 7 ngày
│               ├── documents.js # Upload File .md & Import Google Sheets
│               ├── shopee.js    # Shopee CRUD & Sync Google Sheet
│               ├── customers.js # Lịch sử chat, Bộ lọc, Export CSV, Notes & Tags
│               ├── learning.js  # Learning Queue & AI Auto-Suggest FAQ
│               ├── n8n.js       # Quản lý Workflows n8n & Đồng bộ
│               ├── test.js      # Công cụ Test Bot ngữ nghĩa
│               ├── reports.js   # Báo cáo điều hành AI
│               └── settings.js  # Cấu hình API Keys trực quan
└── knowledge/
    ├── shopee_catalog.json      # Cơ sở dữ liệu sản phẩm Shopee Mall
    ├── zeo_faq.md               # Tài liệu kiến thức chuẩn ZeO
    └── cfc_faq.md               # Tài liệu kiến thức chuẩn CFC Cò Bay"""
p_tree.add_run(tree_text).font.name = 'Courier New'
p_tree.paragraph_format.space_after = Pt(12)

# --- PHẦN 4 ---
add_h1("4. HƯỚNG DẪN VẬN HÀNH & CÁC LỆNH CẦN THIẾT")
add_bullet("cd /Users/hyden/Documents/David-nguyen/N8n/ChatbotN8n/javis/server && source .venv/bin/activate && uvicorn main:app --host 0.0.0.0 --port 8000 --reload", "Khởi động Server:")
add_bullet("Mở trình duyệt truy cập: http://localhost:8000/admin (hoặc http://127.0.0.1:8000/)", "Truy cập Giao Diện Admin:")
add_bullet("Vào tab 'Cài Đặt & API' trên menu để nhập Telegram Bot Token, Chat ID, Google Sheets URL và Google Gemini API Key rồi bấm 'Lưu Tất Cả Cài Đặt'.", "Cấu hình API Keys:")

output_path = Path("/Users/hyden/Documents/David-nguyen/N8n/ChatbotN8n/javis/TAI_LIEU_HE_THONG_CFC_AI.docx")
doc.save(str(output_path))
print(f"Generated DOCX at {output_path}")
